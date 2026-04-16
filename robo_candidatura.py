"""
Robo_CV – Motor principal consolidado

Responsabilidades deste módulo:
- Conectar na OpenAI (via SDK oficial).
- Ler CV (DOCX/PDF), extrair texto e construir um perfil estruturado.
- Analisar uma vaga e extrair JSON com requisitos / ATS.
- Selecionar experiências mais relevantes do perfil.
- Gerar um plano de CV (cv_plan) independente de template.
- Renderizar um DOCX escaneável por ATS.
- Opcionalmente refinar o texto (resumo + bullets) com IA.
- Gerar Cover Letter em DOCX.
- Gerar descrição ATS / LinkedIn.
- Gerar mensagem para recrutador.
- Gerar Guia de Entrevista em DOCX.

Há dois modos:
- MODO INTERATIVO (console/Jupyter): faz perguntas com input().
- MODO SILENCIOSO (API): não usa input(), só IA + texto do CV.
"""

from __future__ import annotations

import os
import re
import json
from datetime import date
from pathlib import Path
from typing import List, Tuple

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except ImportError as exc:
    raise ImportError(
        "pdfminer.six não está instalado corretamente. "
        "Instale com: pip install pdfminer.six"
    ) from exc


# =========================================================
# CONFIGURAÇÃO BÁSICA
# =========================================================

LIMITE_MINIMO_TEXTO_INICIAL = 800
ALLOWED_CV_EXTENSIONS = {".docx", ".pdf"}

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output"
UPLOAD_DIR = BASE_DIR / "api_uploads"

OUTPUT_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise RuntimeError(
        "OPENAI_API_KEY não encontrada no ambiente. "
        "Defina no .env ou nas variáveis de sistema."
    )

client = OpenAI(api_key=api_key)


# =========================================================
# UTILITÁRIOS GERAIS
# =========================================================

def slugificar(texto: str) -> str:
    texto = (texto or "").strip()
    texto = re.sub(r"\s+", " ", texto)
    texto = re.sub(r"[^\w\s-]", "", texto, flags=re.UNICODE)
    return texto.replace(" ", "_") or "arquivo"


def normalizar_texto(texto: str) -> str:
    if not texto:
        return ""
    linhas = [linha.strip() for linha in texto.splitlines()]
    texto_junto = " ".join(linhas)
    return " ".join(texto_junto.split())


def safe_str(value: object, default: str = "") -> str:
    if value is None:
        return default
    return str(value).strip()


def _lower_or_default(value: object, default: str = "") -> str:
    texto = safe_str(value, default)
    return texto.lower() if texto else default


def validar_arquivo_cv(caminho_arquivo: str | Path) -> Path:
    caminho = Path(caminho_arquivo)
    if not caminho.is_file():
        raise FileNotFoundError(f"Arquivo não encontrado: {caminho}")
    if caminho.suffix.lower() not in ALLOWED_CV_EXTENSIONS:
        raise ValueError("arquivocv deve ser .docx ou .pdf.")
    return caminho


def _extrair_json_de_resposta(conteudo: str, mensagem_erro: str) -> dict:
    """
    Tenta fazer json.loads; se falhar, tenta extrair um bloco { ... } do texto.
    Lança ValueError com mensagem clara se nada válido for encontrado.
    """
    try:
        return json.loads(conteudo)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", conteudo, re.DOTALL)
        if not match:
            raise ValueError(mensagem_erro)
        return json.loads(match.group())


# =========================================================
# 1) ANÁLISE DA VAGA
# =========================================================

def analisar_vaga(texto_vaga: str) -> dict:
    system_prompt = """
    Você é um especialista em Recursos Humanos, ATS e Inteligência Artificial.
    Sua tarefa é extrair as informações essenciais da vaga e retornar um JSON LIMPO,
    sem comentários, sem texto fora do JSON.

    O JSON deve conter:
    - titulo_vaga
    - senioridade
    - area
    - local
    - modelo_trabalho
    - responsabilidades
    - requisitos_obrigatorios
    - requisitos_desejaveis
    - skills_tecnicas
    - competencias
    - palavras_chave_ATS
    """

    user_prompt = f"Analise a vaga abaixo e retorne APENAS o JSON:\n\n{texto_vaga}"

    resp = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )

    conteudo = resp.choices[0].message.content
    vaga = _extrair_json_de_resposta(
        conteudo,
        "A IA não retornou um JSON válido para a vaga.",
    )

    vaga.setdefault("titulo_vaga", vaga.get("titulo_vaga") or "Cargo")
    vaga.setdefault("senioridade", vaga.get("senioridade") or "Profissional")
    vaga.setdefault("area", vaga.get("area") or "área")
    vaga.setdefault("local", vaga.get("local") or "")
    vaga.setdefault("modelo_trabalho", vaga.get("modelo_trabalho") or "")
    vaga.setdefault("responsabilidades", vaga.get("responsabilidades") or [])
    vaga.setdefault("requisitos_obrigatorios", vaga.get("requisitos_obrigatorios") or [])
    vaga.setdefault("requisitos_desejaveis", vaga.get("requisitos_desejaveis") or [])
    vaga.setdefault("skills_tecnicas", vaga.get("skills_tecnicas") or [])
    vaga.setdefault("competencias", vaga.get("competencias") or [])
    vaga.setdefault("palavras_chave_ATS", vaga.get("palavras_chave_ATS") or [])
    return vaga


# =========================================================
# 2) EXTRAÇÃO DE TEXTO DO CV (DOCX/PDF)
# =========================================================

def ler_texto_de_docx(caminho: str | Path) -> str:
    doc = Document(str(caminho))
    paragrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return normalizar_texto("\n".join(paragrafos))


def ler_texto_de_pdf(caminho_pdf: str | Path) -> str:
    texto = pdfminer_extract_text(str(caminho_pdf))
    texto_limpo = normalizar_texto(texto)
    if len(texto_limpo) < 100:
        raise ValueError("PDF parece vazio ou ilegível (talvez escaneado).")
    return texto_limpo


# =========================================================
# 3) CONSTRUÇÃO DO PERFIL VIA TEXTO
# =========================================================

def analisar_texto_inicial(texto: str) -> dict:
    texto_limpo = " ".join((texto or "").split())

    if len(texto_limpo) < LIMITE_MINIMO_TEXTO_INICIAL:
        return {
            "status": "incompleto",
            "mensagem": (
                "Seu texto ainda está muito curto para construir um perfil profissional completo.\n\n"
                "Por favor escreva com mais detalhes sobre:\n"
                "- Cada empresa na qual trabalhou\n"
                "- Suas funções e responsabilidades\n"
                "- Projetos realizados\n"
                "- Tecnologias utilizadas\n"
                "- Conquistas mensuráveis (resultados)\n"
                "- Formação acadêmica\n"
                "- Certificações e cursos\n"
                "- Idiomas\n"
                "- Interesses pessoais\n\n"
                "Se algo realmente não existir, escreva não possuo.\n\n"
                f"Caracteres atuais: {len(texto_limpo)} / "
                f"{LIMITE_MINIMO_TEXTO_INICIAL} necessários."
            ),
        }

    return {"status": "ok", "mensagem": ""}


def extrair_perfil_com_IA(texto: str) -> dict:
    system_prompt = """
    Você é um especialista em análise de currículo, RH e extração estruturada.
    Receberá um texto do candidato e deve extrair um PERFIL COMPLETO no formato JSON.

    O JSON deve conter:
    nome
    titulo_padrao
    cidade
    modelo_trabalho_preferido
    telefone
    email
    nacionalidade
    linkedin
    github
    portfolio_url

    formacao: lista
    certificacoes: lista
    competencias_chave: lista
    skills_gerais: lista
    idiomas: lista
    hobbies: lista

    experiencias: lista de objetos:
      - empresa
      - cargo
      - periodo
      - bullets
      - skills

    IMPORTANTE:
    - Se alguma informação não existir no texto, deixe vazio, mas NÃO invente nada.
    - Não adicione empresas, cargos, cursos ou datas que não estejam no texto.
    - Apenas extraia, não reescreva o texto.
    - Responda apenas o JSON, sem explicações.
    """

    user_prompt = f'TEXTO DO CANDIDATO:\n"""\n{texto}\n"""'

    resp = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.0,
    )

    conteudo = resp.choices[0].message.content
    perfil = _extrair_json_de_resposta(
        conteudo,
        "A IA não retornou JSON válido ao extrair o perfil.",
    )
    return perfil


def _normalizar_perfil(perfil: dict) -> dict:
    chaves_lista = [
        "formacao",
        "certificacoes",
        "competencias_chave",
        "skills_gerais",
        "idiomas",
        "hobbies",
        "experiencias",
    ]
    chaves_scalar = [
        "nome",
        "titulo_padrao",
        "cidade",
        "modelo_trabalho_preferido",
        "telefone",
        "email",
        "nacionalidade",
        "linkedin",
        "github",
        "portfolio_url",
    ]
    for chave in chaves_lista:
        perfil.setdefault(chave, [])
    for chave in chaves_scalar:
        perfil.setdefault(chave, "")
    return perfil


def normalizar_perfil(perfil: dict) -> dict:
    return _normalizar_perfil(perfil)


# ---------- MODO SILENCIOSO ----------

def construir_perfil_sem_interacao(texto_inicial: str) -> dict:
    analise = analisar_texto_inicial(texto_inicial)
    perfil = extrair_perfil_com_IA(texto_inicial)
    perfil = normalizar_perfil(perfil)

    if analise["status"] == "incompleto":
        perfil["aviso_texto_curto"] = analise["mensagem"]

    return perfil


def construir_perfil_a_partir_de_arquivo_cv_sem_interacao(
    caminho_arquivo: str | Path,
) -> dict:
    caminho_arquivo = validar_arquivo_cv(caminho_arquivo)
    ext = caminho_arquivo.suffix.lower()

    if ext == ".docx":
        texto = ler_texto_de_docx(caminho_arquivo)
    elif ext == ".pdf":
        texto = ler_texto_de_pdf(caminho_arquivo)
    else:
        raise ValueError(f"Extensão de arquivo não suportada: {ext}")

    return construir_perfil_sem_interacao(texto)


# ---------- MODO INTERATIVO ----------

def detectar_campos_faltantes(perfil: dict) -> List[str]:
    perguntas: List[str] = []

    if not perfil.get("nome"):
        perguntas.append("Qual é o seu nome completo?")

    if not perfil.get("titulo_padrao"):
        perguntas.append(
            "Qual título profissional melhor te representa? Ex: Analista de Dados, Engenheiro de Processos."
        )

    if not perfil.get("cidade"):
        perguntas.append("Em qual cidade/estado você reside atualmente?")

    if not perfil.get("formacao"):
        perguntas.append(
            "Você não descreveu sua formação acadêmica. Qual o curso, instituição e ano de formação? Se não possuir, escreva não possuo."
        )

    if not perfil.get("experiencias"):
        perguntas.append(
            "Você não descreveu experiências profissionais. Pode descrever pelo menos uma experiência relevante? Se não possuir, escreva não possuo."
        )
    else:
        for exp in perfil["experiencias"]:
            if not exp.get("periodo"):
                perguntas.append(
                    f"Qual foi o período em que você trabalhou na empresa {exp.get('empresa', '')}?"
                )
            if not exp.get("bullets"):
                perguntas.append(
                    f"Descreva 3 a 5 responsabilidades/conquistas no cargo {exp.get('cargo', '')} na empresa {exp.get('empresa', '')}."
                )

    if not perfil.get("idiomas"):
        perguntas.append(
            "Quais idiomas você fala e em qual nível? Ex: Inglês Avançado. Se não tiver, escreva não possuo."
        )

    if not perfil.get("certificacoes"):
        perguntas.append(
            "Você possui cursos ou certificações relevantes? Se não, escreva não possuo."
        )

    return perguntas


def atualizar_perfil_com_resposta(perfil: dict, pergunta: str, resposta: str) -> None:
    resp = (resposta or "").strip()
    lower = resp.lower()
    pergunta_lower = pergunta.lower()

    if lower in {"não possuo", "nao possuo", "não tenho", "nao tenho", "nenhum", "nenhuma"}:
        if "formação acadêmica" in pergunta_lower:
            perfil["formacao"] = []
        elif "experiências profissionais" in pergunta_lower:
            perfil["experiencias"] = []
        elif "idiomas" in pergunta_lower:
            perfil["idiomas"] = []
        elif "certificações" in pergunta_lower:
            perfil["certificacoes"] = []
        return

    if "seu nome completo" in pergunta_lower:
        perfil["nome"] = resp
        return

    if "título profissional melhor te representa" in pergunta_lower:
        perfil["titulo_padrao"] = resp
        return

    if "cidade/estado" in pergunta_lower:
        perfil["cidade"] = resp
        return

    if "formação acadêmica" in pergunta_lower:
        atual = perfil.get("formacao", [])
        atual.append(resp)
        perfil["formacao"] = atual
        return

    if "quais idiomas você fala" in pergunta_lower:
        perfil["idiomas"] = [linha.strip() for linha in resp.split(",") if linha.strip()]
        return

    if "cursos ou certificações" in pergunta_lower:
        perfil["certificacoes"] = [linha.strip() for linha in resp.split(",") if linha.strip()]
        return

    if "qual foi o período em que você trabalhou na empresa" in pergunta_lower:
        for exp in perfil.get("experiencias", []):
            empresa = exp.get("empresa", "")
            if empresa and empresa in pergunta:
                exp["periodo"] = resp
                return

    if "descreva 3 a 5 responsabilidades" in pergunta_lower:
        for exp in perfil.get("experiencias", []):
            cargo = exp.get("cargo", "")
            if cargo and cargo in pergunta:
                bullets = [linha.strip("- ").strip() for linha in resp.split("\n") if linha.strip()]
                exp["bullets"] = bullets
                return

    comp = perfil.get("competencias_chave", [])
    comp.append(resp)
    perfil["competencias_chave"] = comp


def construir_perfil_interativo(texto_inicial: str) -> dict:
    analise = analisar_texto_inicial(texto_inicial)
    if analise["status"] == "incompleto":
        raise ValueError(
            "Texto inicial muito curto. Por favor complemente o texto e rode novamente."
        )

    perfil = extrair_perfil_com_IA(texto_inicial)
    perfil = normalizar_perfil(perfil)

    while True:
        perguntas = detectar_campos_faltantes(perfil)
        if not perguntas:
            break

        for pergunta in perguntas:
            resposta = input(f"{pergunta}\n> ")
            atualizar_perfil_com_resposta(perfil, pergunta, resposta)

    return perfil


def construir_perfil_a_partir_de_arquivo_cv_interativo(
    caminho_arquivo: str | Path,
) -> dict:
    caminho_arquivo = validar_arquivo_cv(caminho_arquivo)
    ext = caminho_arquivo.suffix.lower()

    if ext == ".docx":
        texto = ler_texto_de_docx(caminho_arquivo)
    elif ext == ".pdf":
        texto = ler_texto_de_pdf(caminho_arquivo)
    else:
        raise ValueError(f"Extensão de arquivo não suportada: {ext}")

    return construir_perfil_interativo(texto)


# =========================================================
# 4) SELEÇÃO DE EXPERIÊNCIAS E CV_PLAN
# =========================================================

TITULOS_SECOES = {
    "pt": {
        "summary": "Resumo",
        "experience": "Experiência Profissional",
        "education": "Formação Acadêmica",
        "courses": "Cursos & Certificações",
        "skills": "Competências Técnicas",
        "languages": "Idiomas",
        "hobbies": "Interesses Pessoais",
    },
    "en": {
        "summary": "Summary",
        "experience": "Professional Experience",
        "education": "Education",
        "courses": "Courses & Certifications",
        "skills": "Technical Skills",
        "languages": "Languages",
        "hobbies": "Personal Interests",
    },
}


def selecionar_experiencias(vaga: dict, perfil: dict) -> List[dict]:
    experiencias = perfil.get("experiencias", []) or []
    if not experiencias:
        return []

    skills_vaga = set((vaga.get("skills_tecnicas", []) or []) + (vaga.get("palavras_chave_ATS", []) or []))

    pontuadas: List[tuple[int, dict]] = []
    for exp in experiencias:
        score = 0

        for skill in exp.get("skills", []) or []:
            if skill in skills_vaga:
                score += 3

        bullets_texto = " ".join(exp.get("bullets", []) or []).lower()
        for skill in skills_vaga:
            if safe_str(skill).lower() in bullets_texto:
                score += 1

        pontuadas.append((score, exp))

    pontuadas.sort(key=lambda x: x[0], reverse=True)

    total = len(pontuadas)
    if total <= 3:
        max_exp = total
    elif total <= 6:
        max_exp = min(5, total)
    else:
        max_exp = min(8, total)

    return [exp for score, exp in pontuadas[:max_exp]]


def gerar_cv_plan_basico(
    vaga: dict,
    perfil: dict,
    experiencias_relevantes: List[dict],
    idioma: str = "pt",
) -> dict:
    lang = TITULOS_SECOES.get(idioma, TITULOS_SECOES["pt"])
    titulo_vaga = safe_str(vaga.get("titulo_vaga"), safe_str(perfil.get("titulo_padrao"), "Cargo"))
    senioridade = _lower_or_default(vaga.get("senioridade"), "profissional")
    area = _lower_or_default(vaga.get("area"), "área")

    header_links = []
    if safe_str(perfil.get("linkedin")):
        header_links.append(f"LinkedIn: {safe_str(perfil.get('linkedin'))}")
    if safe_str(perfil.get("github")):
        header_links.append(f"GitHub: {safe_str(perfil.get('github'))}")
    if safe_str(perfil.get("portfolio_url")):
        header_links.append(f"Portfólio: {safe_str(perfil.get('portfolio_url'))}")

    header = {
        "name": safe_str(perfil.get("nome")),
        "title": titulo_vaga,
        "location_line": " | ".join(
            [x for x in [safe_str(perfil.get("cidade")), safe_str(perfil.get("modelo_trabalho_preferido"))] if x]
        ),
        "contacts": [
            f"Telefone/WhatsApp: {safe_str(perfil.get('telefone'))}" if safe_str(perfil.get("telefone")) else "",
            f"Email: {safe_str(perfil.get('email'))}" if safe_str(perfil.get("email")) else "",
        ],
        "links": header_links,
    }

    empresas = ", ".join(
        sorted({safe_str(exp.get("empresa")) for exp in (experiencias_relevantes or []) if safe_str(exp.get("empresa"))})
    )
    if not empresas:
        empresas = "relevant companies" if idioma == "en" else "empresas relevantes"

    if idioma == "en":
        resumo_paragrafos = [
            f"{senioridade.capitalize()} professional with experience in {area} and high-performance environments, connecting business needs to analytical solutions and continuous improvement.",
            f"Experience in companies such as {empresas}, focusing on data, automation, operational efficiency and data-driven decision-making.",
        ]
    else:
        resumo_paragrafos = [
            f"Profissional {senioridade} com atuação em {area} e ambientes de alta performance, conectando necessidades de negócio a soluções analíticas e melhoria contínua.",
            f"Experiência em empresas como {empresas}, com foco em dados, automação, eficiência operacional e tomada de decisão orientada por indicadores.",
        ]

    experiencia_items = []
    for exp in experiencias_relevantes or []:
        experiencia_items.append(
            {
                "company": safe_str(exp.get("empresa")),
                "role": safe_str(exp.get("cargo")),
                "period": safe_str(exp.get("periodo")),
                "location": safe_str(perfil.get("cidade")),
                "bullets": exp.get("bullets") or [],
            }
        )

    formacao_items = perfil.get("formacao") or []
    cursos_items = perfil.get("certificacoes") or []
    skills_gerais = perfil.get("skills_gerais") or []
    skills_vaga = vaga.get("skills_tecnicas") or []
    skills_items = []
    if skills_gerais or skills_vaga:
        skills_items = [", ".join(sorted(set(skills_gerais) | set(skills_vaga)))]
    idiomas_items = perfil.get("idiomas") or []
    hobbies_items = perfil.get("hobbies") or []

    sections: List[dict] = []
    sections.append({"title": lang["summary"], "type": "summary", "paragraphs": resumo_paragrafos})

    if experiencia_items:
        sections.append({"title": lang["experience"], "type": "experience", "items": experiencia_items})
    if formacao_items:
        sections.append({"title": lang["education"], "type": "education", "items": formacao_items})
    if cursos_items:
        sections.append({"title": lang["courses"], "type": "courses", "items": cursos_items})
    if skills_items:
        sections.append({"title": lang["skills"], "type": "skills", "items": skills_items})
    if idiomas_items:
        sections.append({"title": lang["languages"], "type": "languages", "items": idiomas_items})
    if hobbies_items:
        sections.append({"title": lang["hobbies"], "type": "hobbies", "items": hobbies_items})

    return {"header": header, "sections": sections, "language": idioma}


# =========================================================
# 5) RENDERIZAÇÃO DO CV EM DOCX
# =========================================================

def renderizar_cv_docx(
    cv_plan: dict,
    nome_arquivo: str = "CV_Robo_CV.docx",
    pasta: str | Path = OUTPUT_DIR,
) -> str:
    pasta = Path(pasta)
    pasta.mkdir(exist_ok=True, parents=True)
    caminho_arquivo = pasta / nome_arquivo

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)
    estilo_normal.paragraph_format.space_after = Pt(2)

    header = cv_plan.get("header", {})

    nome = (header.get("name") or "").upper()
    if nome:
        p_nome = doc.add_paragraph()
        p_nome.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_nome = p_nome.add_run(nome)
        r_nome.bold = True
        r_nome.font.size = Pt(18)

    titulo = header.get("title", "")
    if titulo:
        p_titulo = doc.add_paragraph()
        r_titulo = p_titulo.add_run(titulo)
        r_titulo.bold = True
        r_titulo.font.size = Pt(12)
        r_titulo.font.color.rgb = RGBColor(0x00, 0x32, 0x64)

    location_line = header.get("location_line", "")
    contatos = [c for c in header.get("contacts", []) if c.strip()]
    linha_contato = " | ".join([location_line] + contatos) if location_line else " | ".join(contatos)
    if linha_contato:
        p_contatos = doc.add_paragraph()
        p_contatos.add_run(linha_contato)

    links = header.get("links", []) or []
    if links:
        p_links = doc.add_paragraph()
        r_links = p_links.add_run(" | ".join(links))
        r_links.font.size = Pt(10)

    p_sep = doc.add_paragraph()
    r_sep = p_sep.add_run("_" * 49)
    r_sep.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)

    doc.add_paragraph()

    for sec in cv_plan.get("sections", []):
        title = sec.get("title", "")
        sec_type = sec.get("type", "")

        if title:
            p_title = doc.add_paragraph()
            p_title.paragraph_format.keep_with_next = True
            r_title = p_title.add_run(title.upper())
            r_title.bold = True
            r_title.font.size = Pt(11)
            r_title.font.color.rgb = RGBColor(0x00, 0x32, 0x64)

        if sec_type == "summary":
            for par in sec.get("paragraphs", []):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(par)

        elif sec_type == "experience":
            for item in sec.get("items", []):
                company = item.get("company", "")
                role = item.get("role", "")
                period = item.get("period", "")
                location = item.get("location", "")

                cabecalho = f"{company} | {role}" if role else company
                if period:
                    cabecalho += f" | {period}"

                p_head = doc.add_paragraph()
                p_head.paragraph_format.keep_with_next = True
                r_head = p_head.add_run(cabecalho)
                r_head.bold = True
                r_head.font.size = Pt(11)

                if location:
                    p_loc = doc.add_paragraph()
                    p_loc.paragraph_format.keep_with_next = True
                    r_loc = p_loc.add_run(location)
                    r_loc.font.size = Pt(10)

                for bullet in item.get("bullets", []):
                    p_bullet = doc.add_paragraph(style="List Bullet")
                    p_bullet.add_run(bullet)

                espaco = doc.add_paragraph()
                espaco.paragraph_format.space_after = Pt(4)

        elif sec_type in {"education", "courses", "skills", "languages", "hobbies"}:
            for linha in sec.get("items", []):
                p_item = doc.add_paragraph()
                p_item.paragraph_format.space_after = Pt(1)
                p_item.add_run(linha)

        else:
            for linha in sec.get("items", []) or sec.get("paragraphs", []):
                p_item = doc.add_paragraph()
                p_item.add_run(linha)

        doc.add_paragraph()

    doc.save(caminho_arquivo)
    return str(caminho_arquivo)


# =========================================================
# 6) REFINO DO CV PLAN COM IA
# =========================================================

def refinar_cv_plan_com_ia(cv_plan: dict, vaga: dict, idioma: str = "pt") -> dict:
    system_prompt = f"""
    Você é um especialista sênior em adequação de currículos à vaga.
    Seu trabalho é reescrever o resumo e os bullets do CV de forma altamente alinhada à vaga enviada,
    SEM inventar fatos.

    REGRAS:
    - NÃO inventar datas, resultados, tecnologias, empresas ou cargos.
    - NÃO alterar o header (nome, contatos, cidade).
    - NÃO adicionar nem remover experiências.
    - NÃO mudar o número de bullets de cada experiência.
    - Manter estrutura do JSON EXATAMENTE igual.

    FOCO:
    - Usar linguagem e palavras-chave da vaga.
    - Destacar habilidades transferíveis, deixando claro que derivam das experiências e competências já presentes no CV.

    Idioma do currículo: {idioma.upper()}.
    Responda SOMENTE com o JSON final.
    """

    user_prompt = f"""
    VAGA JSON:
    {json.dumps(vaga, ensure_ascii=False, indent=2)}

    CV_PLAN ORIGINAL JSON:
    {json.dumps(cv_plan, ensure_ascii=False, indent=2)}
    """

    resp = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.25,
    )

    conteudo = resp.choices[0].message.content
    return _extrair_json_de_resposta(
        conteudo,
        "A IA não retornou um JSON válido ao refinar o cv_plan.",
    )


# =========================================================
# 7) DESCRIÇÃO ATS / LINKEDIN
# =========================================================

def gerar_descricao_profissional_ats(perfil: dict, vaga: dict, idioma: str = "pt") -> str:
    nome = perfil.get("nome", "")
    titulo_padrao = perfil.get("titulo_padrao", "")
    cidade = perfil.get("cidade", "")
    experiencias = perfil.get("experiencias", [])
    skills_gerais = perfil.get("skills_gerais", [])
    idiomas_perfil = perfil.get("idiomas", [])

    titulo_vaga = vaga.get("titulo_vaga", "")
    area_vaga = vaga.get("area", "")
    senioridade = vaga.get("senioridade", "")
    skills_vaga = vaga.get("skills_tecnicas", [])
    palavras_chave_vaga = vaga.get("palavras_chave_ATS", [])

    contexto_perfil = {
        "nome": nome,
        "titulo_padrao": titulo_padrao,
        "cidade": cidade,
        "experiencias": experiencias,
        "skills_gerais": skills_gerais,
        "idiomas": idiomas_perfil,
    }

    contexto_vaga = {
        "titulo_vaga": titulo_vaga,
        "area": area_vaga,
        "senioridade": senioridade,
        "skills_vaga": skills_vaga,
        "palavras_chave_vaga": palavras_chave_vaga,
    }

    if idioma == "en":
        system_prompt = """
        You are an expert career positioning and ATS optimization assistant.

        Your task:
        - Generate a short PROFESSIONAL SUMMARY optimized for ATS and LinkedIn.
        - Target the role described in the job posting.
        - Use real keywords from the job skills and requirements.
        - Highlight results, tools and business impact.
        - Use strong action verbs and technical language.
        - DO NOT invent metrics.
        - If there are no explicit numbers, keep the sentences qualitative.
        - Avoid weak expressions like "responsible for" or "supported in".
        - Tone: confident, strategic, focused on being the solution for the role.

        Output:
        - 3 to 5 short paragraphs.
        - Max 8 lines total.
        """
    else:
        system_prompt = """
        Você é um especialista em posicionamento de carreira e otimização para ATS.

        Sua tarefa:
        - Gerar um RESUMO / DESCRIÇÃO PROFISSIONAL otimizado para ATS e LinkedIn.
        - Mirar diretamente na vaga descrita.
        - Usar as palavras-chave reais da vaga.
        - Destacar resultados, ferramentas e impacto no negócio.
        - Usar verbos de ação e linguagem técnica.
        - NÃO inventar métricas.
        - Se não houver números, mantenha frases qualitativas.
        - Evitar expressões fracas como "responsável por" ou "apoio em".
        - Tom firme, confiante, posicionado como solução para os desafios da vaga.

        Formato:
        - 3 a 5 parágrafos curtos.
        - Cerca de 5 a 8 linhas no total.
        """

    user_prompt = f"""
    A seguir estão os dados do perfil do candidato e da vaga.

    PERFIL DO CANDIDATO:
    {json.dumps(contexto_perfil, ensure_ascii=False, indent=2)}

    VAGA:
    {json.dumps(contexto_vaga, ensure_ascii=False, indent=2)}

    Gere apenas o texto da descrição profissional, sem título, sem bullets e sem explicações adicionais.
    """

    resposta = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.4,
    )

    return resposta.choices[0].message.content.strip()


# =========================================================
# 8) MENSAGEM PARA RECRUTADOR
# =========================================================

def gerar_mensagem_recrutador(perfil: dict, vaga: dict, idioma: str = "pt") -> dict:
    nome = perfil.get("nome", "")
    titulo_padrao = perfil.get("titulo_padrao", "")
    experiencias = perfil.get("experiencias", [])
    skills_gerais = perfil.get("skills_gerais", [])

    empresa_vaga = vaga.get("empresa") or vaga.get("local") or ""
    titulo_vaga = vaga.get("titulo_vaga", "")
    area_vaga = vaga.get("area", "")
    skills_vaga = vaga.get("skills_tecnicas", [])
    palavras_chave_vaga = vaga.get("palavras_chave_ATS", [])

    contexto = {
        "nome": nome,
        "titulo_padrao": titulo_padrao,
        "experiencias": experiencias,
        "skills_gerais": skills_gerais,
        "empresa_vaga": empresa_vaga,
        "titulo_vaga": titulo_vaga,
        "area_vaga": area_vaga,
        "skills_vaga": skills_vaga,
        "palavras_chave_vaga": palavras_chave_vaga,
    }

    if idioma == "en":
        system_prompt = """
        You are an expert recruiter and career coach.

        Task:
        - Create a 3-line outreach message to a recruiter on LinkedIn or via email about a specific role.
        - Structure:
          1) Personalized hook
          2) Connection between the candidate's key skills/results and the role
          3) Light CTA
        - Use real keywords from the job.
        - Use results only if they clearly appear in the profile/experiences.
        - NEVER invent metrics or achievements.
        - Tone: concise, respectful, confident.

        Output:
        - Return a JSON with fields: linha1, linha2, linha3, mensagem_unica
        """
    else:
        system_prompt = """
        Você é um recrutador experiente e coach de carreira.

        Tarefa:
        - Criar uma mensagem de 3 linhas para abordar um recrutador no LinkedIn ou por e-mail sobre uma vaga específica.
        - Estrutura:
          1) Gancho personalizado
          2) Conexão entre skills/resultados do candidato e as necessidades da vaga
          3) CTA leve
        - Usar palavras-chave reais da vaga.
        - Usar resultados apenas se eles aparecerem claramente no perfil/experiências.
        - NUNCA inventar métricas ou conquistas.
        - Tom conciso, respeitoso, confiante.

        Formato:
        - Retorne um JSON com os campos: linha1, linha2, linha3, mensagem_unica
        """

    user_prompt = f"""
    A seguir estão os dados do candidato e da vaga.

    PERFIL DO CANDIDATO:
    {json.dumps(contexto, ensure_ascii=False, indent=2)}

    Gere a mensagem seguindo fielmente o formato JSON pedido.
    Não inclua comentários, explicações ou texto extra fora do JSON.
    """

    resposta = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.4,
    )

    conteudo = resposta.choices[0].message.content.strip()

    try:
        dados = json.loads(conteudo)
    except json.JSONDecodeError:
        dados = {
            "linha1": "",
            "linha2": "",
            "linha3": "",
            "mensagem_unica": conteudo,
        }

    dados.setdefault("linha1", "")
    dados.setdefault("linha2", "")
    dados.setdefault("linha3", "")
    dados.setdefault(
        "mensagem_unica",
        "\n".join([x for x in [dados["linha1"], dados["linha2"], dados["linha3"]] if x]),
    )
    return dados


# =========================================================
# 9) LEITURA DA VAGA (TEXTO / URL / PDF)
# =========================================================

def extrair_texto_de_pagina_web(url: str) -> str:
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    main = soup.find("main")
    texto = main.get_text(separator=" ", strip=True) if main else soup.get_text(separator=" ", strip=True)
    texto_limpo = normalizar_texto(texto)

    return texto_limpo


def extrair_texto_de_pdf_vaga(caminho_pdf: str | Path) -> str:
    texto = ler_texto_de_pdf(caminho_pdf)
    return normalizar_texto(texto)


def carregar_texto_vaga(entrada: str) -> str:
    entrada = (entrada or "").strip()
    if not entrada:
        return ""

    if entrada.lower().startswith(("http://", "https://")):
        try:
            return extrair_texto_de_pagina_web(entrada)
        except Exception:
            return ""

    if entrada.lower().endswith(".pdf") and Path(entrada).is_file():
        try:
            return extrair_texto_de_pdf_vaga(entrada)
        except Exception:
            return ""

    return normalizar_texto(entrada)


# =========================================================
# 10) FUNÇÕES AUXILIARES DE ORQUESTRAÇÃO
# =========================================================

def preparar_insumos_basicos(
    caminho_cv: str | Path,
    texto_vaga: str,
    interativo: bool = False,
) -> Tuple[dict, dict]:
    caminho_cv = validar_arquivo_cv(caminho_cv)
    texto_vaga = normalizar_texto(texto_vaga or "")
    if not texto_vaga:
        raise ValueError("Texto da vaga vazio ou inválido.")

    if interativo:
        perfil = construir_perfil_a_partir_de_arquivo_cv_interativo(caminho_cv)
    else:
        perfil = construir_perfil_a_partir_de_arquivo_cv_sem_interacao(caminho_cv)

    perfil = normalizar_perfil(perfil or {})
    vaga = analisar_vaga(texto_vaga)
    return perfil, vaga


def _hoje_formatado(idioma: str = "pt") -> str:
    hoje = date.today()

    if idioma == "pt":
        meses = [
            "janeiro", "fevereiro", "março", "abril", "maio", "junho",
            "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
        ]
        return f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

    meses_en = [
        "January", "February", "March", "April", "May", "June",
        "July", "August", "September", "October", "November", "December",
    ]
    return f"{meses_en[hoje.month - 1]} {hoje.day}, {hoje.year}"


# =========================================================
# 11) COVER LETTER
# =========================================================

def gerar_cover_letter_com_ia(perfil: dict, vaga: dict, idioma: str = "pt") -> dict:
    system_prompt = f"""
    Você é um ESPECIALISTA EM COVER LETTER / CARTA DE APRESENTAÇÃO.
    Você receberá:
    - PERFIL do candidato
    - VAGA analisada

    Sua missão:
    Gerar uma cover letter curta, objetiva e convincente, com foco em ATS e leitura humana.

    REGRAS OBRIGATÓRIAS:
    1. NUNCA invente empresas, cargos, datas, ferramentas, certificações ou números.
    2. Use SOMENTE fatos que estejam no perfil ou na vaga.
    3. Se faltar informação relevante, escreva de forma QUALITATIVA e honesta.
    4. Tom seguro, direto, profissional.
    5. Idioma: {idioma.upper()}.

    FORMATO DE SAÍDA:
    APENAS JSON com:
    - linha_assunto
    - saudacao
    - paragrafos
    - encerramento
    - assinatura_nome
    """

    payload = {
        "perfil": {
            "nome": perfil.get("nome", ""),
            "titulo_padrao": perfil.get("titulo_padrao", ""),
            "cidade": perfil.get("cidade", ""),
            "email": perfil.get("email", ""),
            "telefone": perfil.get("telefone", ""),
            "linkedin": perfil.get("linkedin", ""),
            "skills_gerais": perfil.get("skills_gerais", []),
            "experiencias": perfil.get("experiencias", []),
            "formacao": perfil.get("formacao", []),
            "certificacoes": perfil.get("certificacoes", []),
            "idiomas": perfil.get("idiomas", []),
        },
        "vaga": vaga,
    }

    resp = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": f"Dados:\n{json.dumps(payload, ensure_ascii=False, indent=2)}"},
        ],
        temperature=0.3,
    )

    conteudo = resp.choices[0].message.content
    plano = _extrair_json_de_resposta(
        conteudo,
        "Cover letter: a IA não retornou um JSON válido.",
    )

    if not isinstance(plano, dict):
        raise ValueError("Cover letter: resposta inválida (não é dict).")

    if "paragrafos" not in plano or not isinstance(plano["paragrafos"], list) or len(plano["paragrafos"]) < 2:
        raise ValueError("Cover letter JSON sem parágrafos adequados.")

    plano.setdefault("linha_assunto", "")
    plano.setdefault("saudacao", "Prezados(as)," if idioma == "pt" else "Dear Hiring Team,")
    plano.setdefault("encerramento", "Atenciosamente," if idioma == "pt" else "Sincerely,")
    plano.setdefault("assinatura_nome", perfil.get("nome", ""))
    return plano


def gerar_docx_cover_letter(
    plano_cover_letter: dict,
    perfil: dict,
    empresa_alvo: str,
    cargo_alvo: str,
    idioma: str = "pt",
    pasta: str | Path = OUTPUT_DIR,
) -> str:
    pasta = Path(pasta)
    pasta.mkdir(exist_ok=True, parents=True)

    nome_candidato = perfil.get("nome", "Candidato")
    slug_nome = slugificar(nome_candidato)
    empresa_slug = slugificar(empresa_alvo or "Empresa")
    cargo_slug = slugificar(cargo_alvo or "Cargo")
    nome_arquivo = f"{slug_nome}_CoverLetter_{empresa_slug}_{cargo_slug}.docx"
    caminho_arquivo = pasta / nome_arquivo

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)
    estilo_normal.paragraph_format.space_after = Pt(2)

    p_nome = doc.add_paragraph()
    r_nome = p_nome.add_run(nome_candidato.upper() if nome_candidato else "")
    r_nome.bold = True
    r_nome.font.size = Pt(14)

    titulo = perfil.get("titulo_padrao", "")
    if titulo:
        p_titulo = doc.add_paragraph()
        r_titulo = p_titulo.add_run(titulo)
        r_titulo.font.size = Pt(11)
        r_titulo.font.color.rgb = RGBColor(0x00, 0x32, 0x64)

    contatos = []
    if perfil.get("email"):
        contatos.append(str(perfil["email"]))
    if perfil.get("telefone"):
        contatos.append(str(perfil["telefone"]))
    if perfil.get("linkedin"):
        contatos.append(str(perfil["linkedin"]))
    if contatos:
        doc.add_paragraph(" | ".join(contatos))

    if cargo_alvo or empresa_alvo:
        doc.add_paragraph(f"{cargo_alvo} | {empresa_alvo}".strip(" | "))

    doc.add_paragraph(_hoje_formatado(idioma))

    doc.add_paragraph()
    doc.add_paragraph(plano_cover_letter.get("saudacao", ""))

    for paragrafo in plano_cover_letter.get("paragrafos", []):
        doc.add_paragraph(paragrafo)

    doc.add_paragraph(plano_cover_letter.get("encerramento", ""))
    doc.add_paragraph(plano_cover_letter.get("assinatura_nome", nome_candidato))

    doc.save(caminho_arquivo)
    return str(caminho_arquivo)


# =========================================================
# 12) GUIA DE ENTREVISTA
# =========================================================

def gerar_plano_entrevista_com_ia(
    perfil: dict,
    vaga: dict,
    cv_plan: dict,
    idioma: str = "pt",
) -> dict:
    system_prompt = f"""
    Você é um ESPECIALISTA EM ENTREVISTAS DE EMPREGO, TREINAMENTO DE CANDIDATOS
    e PREPARAÇÃO PARA RECRUTADORES.

    Você receberá:
    - PERFIL estruturado do candidato
    - VAGA analisada
    - CV_PLAN

    Sua missão:
    gerar um GUIA COMPLETO DE ENTREVISTA ajudando o candidato a:
    1) Responder bem às perguntas mais prováveis
    2) Explicar e vender os bullets do currículo
    3) Conduzir uma conversa curta com recrutador

    REGRAS:
    - NUNCA inventar fatos.
    - Quando faltar informação, marcar confiança baixa.
    - Gerar entre 8 e 15 perguntas equilibrando geral, técnica e comportamental.
    - Idioma: {idioma.upper()}.

    FORMATO DE SAÍDA:
    JSON com:
    - perguntas_entrevista
    - orientacoes_bullets
    - roteiro_recrutador
    """

    user_prompt = f"""
    PERFIL:
    {json.dumps(perfil, ensure_ascii=False, indent=2)}

    VAGA:
    {json.dumps(vaga, ensure_ascii=False, indent=2)}

    CV_PLAN:
    {json.dumps(cv_plan, ensure_ascii=False, indent=2)}
    """

    resp = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
    )

    conteudo = resp.choices[0].message.content
    plano = _extrair_json_de_resposta(
        conteudo,
        "A IA não retornou um JSON válido para o plano de entrevista.",
    )

    plano.setdefault("perguntas_entrevista", [])
    plano.setdefault("orientacoes_bullets", [])
    plano.setdefault("roteiro_recrutador", {})
    return plano


def gerar_docx_guia_entrevista(
    plano_entrevista: dict,
    nome_candidato: str,
    empresa_alvo: str,
    cargo_alvo: str,
    idioma: str = "pt",
    pasta: str | Path = OUTPUT_DIR,
) -> str:
    pasta = Path(pasta)
    pasta.mkdir(exist_ok=True, parents=True)

    slug_nome = slugificar(nome_candidato or "Candidato")
    empresa_slug = slugificar(empresa_alvo or "Empresa")
    cargo_slug = slugificar(cargo_alvo or "Cargo")
    nome_arquivo = f"{slug_nome}_GuiaEntrevista_{empresa_slug}_{cargo_slug}.docx"
    caminho_arquivo = pasta / nome_arquivo

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)
    estilo_normal.paragraph_format.space_after = Pt(2)

    titulo = "Guia de Entrevista" if idioma == "pt" else "Interview Guide"
    subtitulo = f"Vaga: {cargo_alvo} | {empresa_alvo}".strip(" | ")

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_titulo = p_titulo.add_run(titulo)
    r_titulo.bold = True
    r_titulo.font.size = Pt(16)
    r_titulo.font.color.rgb = RGBColor(0x00, 0x32, 0x64)

    if subtitulo:
        doc.add_paragraph(subtitulo)

    intro_pt = (
        "Este guia foi gerado automaticamente com base no seu currículo e na vaga. "
        "Ele traz perguntas prováveis de entrevista, respostas sugeridas e orientações "
        "para você conseguir vender suas experiências com segurança."
    )
    intro_en = (
        "This interview guide was automatically generated based on your CV and the job posting. "
        "It includes likely interview questions, suggested answers and guidance to help you "
        "present your experience clearly and confidently."
    )
    doc.add_paragraph(intro_pt if idioma == "pt" else intro_en)

    def add_sec_title(texto: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x00, 0x32, 0x64)

    perguntas = plano_entrevista.get("perguntas_entrevista", []) or []
    if perguntas:
        add_sec_title("Perguntas de entrevista" if idioma == "pt" else "Interview questions")
        for idx, item in enumerate(perguntas, start=1):
            pergunta = item.get("pergunta", "")
            base_resposta = item.get("base_resposta", "")
            resposta_sugerida = item.get("resposta_sugerida", "")
            nota_origem = item.get("nota_origem", "")
            confianca = item.get("confianca", "")
            alerta = item.get("alerta", "")

            p = doc.add_paragraph()
            r = p.add_run(f"{idx}. {pergunta}")
            r.bold = True

            if base_resposta:
                doc.add_paragraph(f"Base: {base_resposta}")
            if resposta_sugerida:
                doc.add_paragraph(f"Resposta sugerida: {resposta_sugerida}")
            if nota_origem:
                doc.add_paragraph(f"Origem: {nota_origem}")
            if confianca:
                doc.add_paragraph(f"Confiança: {confianca}")
            if alerta and alerta.lower() != "nenhum":
                p_alerta = doc.add_paragraph()
                r_alerta = p_alerta.add_run(f"Alerta: {alerta}")
                r_alerta.italic = True
                r_alerta.font.size = Pt(9)

    bullets = plano_entrevista.get("orientacoes_bullets", []) or []
    if bullets:
        add_sec_title("Como explicar os bullets do seu CV" if idioma == "pt" else "How to explain your CV bullets")
        for item in bullets:
            bullet = item.get("bullet_cv", "")
            como_explicar = item.get("como_explicar", "")
            obs = item.get("observacao", "")

            if bullet:
                p_b = doc.add_paragraph()
                r_b = p_b.add_run(f"- {bullet}")
                r_b.bold = True
            if como_explicar:
                doc.add_paragraph(como_explicar)
            if obs:
                p_obs = doc.add_paragraph()
                r_obs = p_obs.add_run(obs)
                r_obs.italic = True
                r_obs.font.size = Pt(9)

    roteiro = plano_entrevista.get("roteiro_recrutador", {}) or {}
    if any((roteiro.get(k) or "").strip() for k in ("abertura", "meio", "fechamento")):
        add_sec_title("Roteiro de conversa (15 min)" if idioma == "pt" else "Recruiter call script (15 min)")
        if roteiro.get("abertura"):
            p = doc.add_paragraph()
            r = p.add_run("Abertura: ")
            r.bold = True
            p.add_run(roteiro["abertura"])
        if roteiro.get("meio"):
            p = doc.add_paragraph()
            r = p.add_run("Meio: ")
            r.bold = True
            p.add_run(roteiro["meio"])
        if roteiro.get("fechamento"):
            p = doc.add_paragraph()
            r = p.add_run("Fechamento: ")
            r.bold = True
            p.add_run(roteiro["fechamento"])

    doc.save(caminho_arquivo)
    return str(caminho_arquivo)


# =========================================================
# 13) FUNÇÕES END-TO-END
# =========================================================

def gerar_cv_end_to_end(
    caminho_cv: str | Path,
    texto_vaga: str,
    empresa_alvo: str,
    cargo_alvo: str,
    idioma: str = "pt",
    refinar_texto: bool = False,
    interativo: bool = False,
) -> str:
    perfil, vaga = preparar_insumos_basicos(caminho_cv, texto_vaga, interativo=interativo)
    experiencias = selecionar_experiencias(vaga, perfil)
    cv_plan = gerar_cv_plan_basico(vaga, perfil, experiencias, idioma=idioma)

    if refinar_texto:
        cv_plan = refinar_cv_plan_com_ia(cv_plan, vaga, idioma=idioma)

    nome_base = slugificar(perfil.get("nome", "Candidato"))
    empresa_slug = slugificar(empresa_alvo or "Empresa")
    cargo_slug = slugificar(cargo_alvo or vaga.get("titulo_vaga", "Cargo"))
    nome_arquivo = f"{nome_base}_CV_{empresa_slug}_{cargo_slug}.docx"

    return renderizar_cv_docx(cv_plan, nome_arquivo=nome_arquivo, pasta=OUTPUT_DIR)


def gerar_cover_letter_end_to_end(
    caminho_cv: str | Path,
    texto_vaga: str,
    empresa_alvo: str,
    cargo_alvo: str,
    idioma: str = "pt",
    interativo: bool = False,
) -> str:
    perfil, vaga = preparar_insumos_basicos(caminho_cv, texto_vaga, interativo=interativo)
    plano = gerar_cover_letter_com_ia(perfil, vaga, idioma=idioma)
    return gerar_docx_cover_letter(
        plano_cover_letter=plano,
        perfil=perfil,
        empresa_alvo=empresa_alvo,
        cargo_alvo=cargo_alvo,
        idioma=idioma,
        pasta=OUTPUT_DIR,
    )


def gerar_guia_entrevista_end_to_end(
    caminho_cv: str | Path,
    texto_vaga: str,
    empresa_alvo: str,
    cargo_alvo: str,
    idioma: str = "pt",
    refinar_texto: bool = False,
) -> str:
    perfil, vaga = preparar_insumos_basicos(caminho_cv, texto_vaga, interativo=False)
    experiencias = selecionar_experiencias(vaga, perfil)
    cv_plan = gerar_cv_plan_basico(vaga, perfil, experiencias, idioma=idioma)

    if refinar_texto:
        cv_plan = refinar_cv_plan_com_ia(cv_plan, vaga, idioma=idioma)

    plano_entrevista = gerar_plano_entrevista_com_ia(
        perfil=perfil,
        vaga=vaga,
        cv_plan=cv_plan,
        idioma=idioma,
    )

    return gerar_docx_guia_entrevista(
        plano_entrevista=plano_entrevista,
        nome_candidato=perfil.get("nome", "Candidato"),
        empresa_alvo=empresa_alvo,
        cargo_alvo=cargo_alvo,
        idioma=idioma,
        pasta=OUTPUT_DIR,
    )


# =========================================================
# 14) COMPATIBILIDADE COM NOMES LEGADOS (OPCIONAL)
# =========================================================

# Mantidos apenas para não quebrar código legado, concentrados no final.
normalizarperfil = normalizar_perfil
carregartextovaga = carregar_texto_vaga
gerarcvendtoend = gerar_cv_end_to_end
gerarcoverletterendtoend = gerar_cover_letter_end_to_end
gerarguiaentrevistaendtoend = gerar_guia_entrevista_end_to_end
gerardescricaoprofissionalats = gerar_descricao_profissional_ats
gerarmensagemrecrutador = gerar_mensagem_recrutador